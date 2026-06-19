from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DIAGRAM_DIR = ROOT / "output" / "doc" / "qa_v16" / "new_diagrams"
MICU_DIR = ROOT / "output" / "doc" / "qa_v16" / "micu"


FIGURES = [
    {
        "before_heading": "1.1.2 大语言模型赋能业务分析的意义",
        "image": MICU_DIR / "cold_chain_context_micu.png",
        "width": 4.85,
        "cn": "图 1-1 冷链仓储安全管理研究背景图",
        "en": "Figure 1-1 Research background diagram of cold-chain warehousing safety management",
    },
    {
        "before_heading": "1.4 论文组织结构",
        "image": DIAGRAM_DIR / "fig_1_2_research_contents.png",
        "width": 6.45,
        "cn": "图 1-2 论文研究内容结构图",
        "en": "Figure 1-2 Structure diagram of the thesis research contents",
    },
    {
        "before_heading": "2.1 Spring Boot 与后端分层开发",
        "image": DIAGRAM_DIR / "fig_2_1_key_technology_relation.png",
        "width": 6.45,
        "cn": "图 2-1 系统关键技术关系图",
        "en": "Figure 2-1 Relationship diagram of key technologies of the system",
    },
    {
        "before_heading": "4 系统概要设计",
        "image": DIAGRAM_DIR / "fig_3_5_requirement_design_mapping.png",
        "width": 6.45,
        "cn": "图 3-5 需求分析与设计映射关系图",
        "en": "Figure 3-5 Mapping relationship diagram between requirements analysis and design",
    },
]


def latest_v15() -> Path:
    desktop = Path.home() / "Desktop"
    matches = list(desktop.glob("**/*v15.docx"))
    if not matches:
        raise FileNotFoundError("Could not find v15 docx under Desktop")
    return max(matches, key=lambda p: p.stat().st_mtime)


def set_run_font(run, name: str, size_pt: float, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def format_caption_para(para) -> None:
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.0


def add_chinese_caption(para, text: str) -> None:
    # Keep the figure number visually close to the existing v15 captions.
    parts = text.split(" ", 2)
    if len(parts) == 3:
        prefix, number, rest = parts
        r1 = para.add_run(prefix + " ")
        set_run_font(r1, "楷体", 10.5)
        r2 = para.add_run(number)
        r2.font.size = Pt(10.5)
        r3 = para.add_run(" " + rest)
        set_run_font(r3, "楷体", 10.5)
    else:
        r = para.add_run(text)
        set_run_font(r, "楷体", 10.5)


def add_english_caption(para, text: str) -> None:
    r = para.add_run(text)
    set_run_font(r, "Times New Roman", 10.5)


def find_paragraph(doc: Document, text: str):
    for para in doc.paragraphs:
        if para.text.strip() == text:
            return para
    raise ValueError(f"Target paragraph not found: {text}")


def insert_figure_before(target_para, fig: dict[str, object]) -> None:
    image_path = Path(fig["image"])
    if not image_path.exists():
        raise FileNotFoundError(str(image_path))

    image_para = target_para.insert_paragraph_before()
    image_para.style = target_para.part.document.styles["Normal"]
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.paragraph_format.space_before = Pt(3)
    image_para.paragraph_format.space_after = Pt(0)
    image_para.paragraph_format.keep_with_next = True
    image_para.add_run().add_picture(str(image_path), width=Inches(float(fig["width"])))

    cn_para = target_para.insert_paragraph_before()
    cn_para.style = target_para.part.document.styles["Normal"]
    format_caption_para(cn_para)
    cn_para.paragraph_format.keep_with_next = True
    add_chinese_caption(cn_para, str(fig["cn"]))

    en_para = target_para.insert_paragraph_before()
    en_para.style = target_para.part.document.styles["Normal"]
    format_caption_para(en_para)
    add_english_caption(en_para, str(fig["en"]))


def main() -> None:
    src = latest_v15()
    dst = src.with_name(src.name.replace("v15", "v16"))
    shutil.copy2(src, dst)
    doc = Document(str(dst))

    for fig in FIGURES:
        target = find_paragraph(doc, str(fig["before_heading"]))
        insert_figure_before(target, fig)

    doc.save(str(dst))
    print(dst)


if __name__ == "__main__":
    main()
