from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn


SRC = Path(r"C:\Users\23869\Desktop\毕设\草稿\郭鑫瑞毕业论文v14.docx")
DST = Path(r"C:\Users\23869\Desktop\毕设\草稿\郭鑫瑞毕业论文v15.docx")


def clear_page_breaks(paragraph) -> None:
    for br in list(paragraph._p.xpath(".//w:br")):
        parent = br.getparent()
        if parent is not None:
            parent.remove(br)


def add_two_page_breaks(paragraph) -> None:
    clear_page_breaks(paragraph)
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
    run.add_break(WD_BREAK.PAGE)


def main() -> None:
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    zh_keyword = None
    en_keyword = None
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("关键词："):
            zh_keyword = paragraph
        if text.startswith("KEY WORDS:"):
            en_keyword = paragraph

    if zh_keyword is None or en_keyword is None:
        raise RuntimeError("keyword paragraph not found")

    # Preserve the intentional reverse-side blank pages after the Chinese and English abstracts.
    add_two_page_breaks(zh_keyword)
    add_two_page_breaks(en_keyword)

    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        from docx.oxml import OxmlElement

        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)

    doc.save(str(DST))
    print(DST)


if __name__ == "__main__":
    main()
