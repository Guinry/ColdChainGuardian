from __future__ import annotations
import copy
import shutil
import zipfile
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

ROOT = Path(r'C:\Users\23869\Desktop\ColdChainGuardian')
V11 = Path(r'C:\Users\23869\Desktop\毕设\草稿\郭鑫瑞毕业论文v11.docx')
V12 = Path(r'C:\Users\23869\Desktop\毕设\草稿\郭鑫瑞毕业论文v12.docx')
V13 = Path(r'C:\Users\23869\Desktop\毕设\草稿\郭鑫瑞毕业论文v13.docx')
STABLE_DIAGRAMS = ROOT / 'output' / 'doc' / 'qa_v11' / 'diagrams_stable'


def clear_para(paragraph):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def clone_runs(dst_p, src_p):
    clear_para(dst_p)
    dst_p.style = src_p.style
    dst_p.alignment = src_p.alignment
    dst_p.paragraph_format.left_indent = src_p.paragraph_format.left_indent
    dst_p.paragraph_format.right_indent = src_p.paragraph_format.right_indent
    dst_p.paragraph_format.first_line_indent = src_p.paragraph_format.first_line_indent
    dst_p.paragraph_format.space_before = src_p.paragraph_format.space_before
    dst_p.paragraph_format.space_after = src_p.paragraph_format.space_after
    dst_p.paragraph_format.line_spacing = src_p.paragraph_format.line_spacing
    for src_run in src_p.runs:
        dst_p._p.append(copy.deepcopy(src_run._r))


def find_heading(doc, text, style='Heading 1'):
    for i, p in enumerate(doc.paragraphs):
        if (p.text or '').strip() == text and (style is None or p.style.name == style):
            return i
    raise ValueError(f'heading not found: {text}')


def body_children(doc):
    return list(doc._body._element)


def element_index(doc, pidx):
    children = body_children(doc)
    return children.index(doc.paragraphs[pidx]._p)


def replace_block_from_doc(dst_doc, dst_start_idx, dst_end_idx, src_doc, src_start_idx, src_end_idx):
    body = dst_doc._body._element
    dst_children = body_children(dst_doc)
    start = element_index(dst_doc, dst_start_idx)
    end = element_index(dst_doc, dst_end_idx)
    anchor = dst_children[end]
    for el in dst_children[start:end]:
        body.remove(el)
    src_children = body_children(src_doc)
    s = element_index(src_doc, src_start_idx)
    e = element_index(src_doc, src_end_idx)
    for el in src_children[s:e]:
        anchor.addprevious(copy.deepcopy(el))


def copy_abstracts(dst_doc, src_doc):
    # Chinese and English abstract paragraphs from v11. Keep v12 cover and intentional blank reverse pages untouched.
    mappings = [
        (22, 23),  # Chinese abstract text block in v12 <- first v11 abstract paragraph
        (23, 24),  # new extra abstract paragraph inserted after mapping below; handled by block replacement instead
    ]
    # Safer: replace paragraphs from 摘 要 heading following text through keywords by cloning v11 equivalent structure.
    def idx_text(doc, text):
        for i,p in enumerate(doc.paragraphs):
            if (p.text or '').strip() == text:
                return i
        raise ValueError(text)
    zh_start_dst = idx_text(dst_doc, '摘  要') + 1
    zh_end_dst = next(i for i in range(zh_start_dst, len(dst_doc.paragraphs)) if dst_doc.paragraphs[i].text.strip().startswith('DESIGN AND IMPLEMENTATION'))
    zh_start_src = idx_text(src_doc, '摘  要') + 1
    zh_end_src = next(i for i in range(zh_start_src, len(src_doc.paragraphs)) if src_doc.paragraphs[i].text.strip().startswith('DESIGN AND IMPLEMENTATION'))
    replace_block_from_doc(dst_doc, zh_start_dst, zh_end_dst, src_doc, zh_start_src, zh_end_src)

    # The replacement invalidates paragraph collection; reopen-like object is not possible, caller reloads after save.


def replace_image_by_caption(doc, caption_prefix, image_path, width_inches=5.45):
    target = None
    for i, p in enumerate(doc.paragraphs):
        if (p.text or '').strip().startswith(caption_prefix):
            target = i
            break
    if target is None:
        raise ValueError(f'caption not found: {caption_prefix}')
    image_p = None
    for j in range(target - 1, max(-1, target - 8), -1):
        if doc.paragraphs[j]._p.xpath('.//a:blip'):
            image_p = doc.paragraphs[j]
            break
    if image_p is None:
        raise ValueError(f'image paragraph not found before {caption_prefix}')
    clear_para(image_p)
    image_p.alignment = 1
    image_p.paragraph_format.space_before = None
    image_p.paragraph_format.space_after = None
    run = image_p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def ensure_update_fields(doc):
    settings = doc.settings.element
    if settings.find(qn('w:updateFields')) is None:
        update_fields = OxmlElement('w:updateFields')
        update_fields.set(qn('w:val'), 'true')
        settings.append(update_fields)


def main():
    shutil.copy2(V12, V13)
    src = Document(str(V11))
    dst = Document(str(V13))

    # 1) Teacher comment P3: abstract replacement. Use v11's corrected abstract/keywords only.
    copy_abstracts(dst, src)
    dst.save(str(V13))

    # Reload after structural replacement of abstracts.
    src = Document(str(V11))
    dst = Document(str(V13))

    # 2) Teacher comments P14-P21: Chapter 1 literature/cuts and Chapter 2 only-introduce-technology.
    replace_block_from_doc(
        dst,
        find_heading(dst, '1 绪论'),
        find_heading(dst, '3 系统分析'),
        src,
        find_heading(src, '1 绪论'),
        find_heading(src, '3 系统分析'),
    )
    dst.save(str(V13))

    # Reload again, then replace comment-related engineering diagrams with larger readable versions.
    dst = Document(str(V13))
    replacements = [
        ('图 3-1 用户角色权限关系图', 'image1_roles_permissions.png', 5.65),
        ('图 3-2 系统业务流程图', 'image2_business_flow.png', 5.70),
        ('图 3-3 告警工单闭环流程图', 'image3_alert_loop.png', 5.65),
        ('图 3-4 系统用例图', 'image4_use_case.png', 5.70),
        ('图 4-1 系统总体架构图', 'image5_architecture.png', 5.70),
        ('图 4-2 系统部署结构图', 'image6_deployment.png', 5.70),
        ('图 4-3 系统功能模块图', 'image7_modules.png', 5.70),
        ('图 4-4 数据库 E-R 图', 'image8_er.png', 5.55),
        ('图 4-5 系统核心类图', 'image9_class.png', 5.55),
    ]
    for cap, fname, width in replacements:
        replace_image_by_caption(dst, cap, STABLE_DIAGRAMS / fname, width)

    ensure_update_fields(dst)
    dst.save(str(V13))
    print(V13)

if __name__ == '__main__':
    main()
