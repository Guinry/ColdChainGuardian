from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:\Users\23869\Desktop\ColdChainGuardian")
V11 = Path(r"C:\Users\23869\Desktop\毕设\草稿\郭鑫瑞毕业论文v11.docx")
V12 = Path(r"C:\Users\23869\Desktop\毕设\草稿\郭鑫瑞毕业论文v12.docx")
V14 = Path(r"C:\Users\23869\Desktop\毕设\草稿\郭鑫瑞毕业论文v14.docx")
STABLE_DIAGRAMS = ROOT / "output" / "doc" / "qa_v11" / "diagrams_stable"


def body_children(doc: Document):
    return list(doc._body._element)


def element_index(doc: Document, paragraph_index: int) -> int:
    return body_children(doc).index(doc.paragraphs[paragraph_index]._p)


def find_heading(doc: Document, text: str, style: str = "Heading 1") -> int:
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == text and paragraph.style.name == style:
            return i
    raise ValueError(f"heading not found: {text}")


def find_text(doc: Document, text: str) -> int:
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == text:
            return i
    raise ValueError(f"text not found: {text}")


def clear_para(paragraph: Paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def copy_paragraph_format(dst: Paragraph, src: Paragraph) -> None:
    dst.style = src.style.name
    dst.alignment = src.alignment
    dst.paragraph_format.left_indent = src.paragraph_format.left_indent
    dst.paragraph_format.right_indent = src.paragraph_format.right_indent
    dst.paragraph_format.first_line_indent = src.paragraph_format.first_line_indent
    dst.paragraph_format.space_before = src.paragraph_format.space_before
    dst.paragraph_format.space_after = src.paragraph_format.space_after
    dst.paragraph_format.line_spacing = src.paragraph_format.line_spacing
    dst.paragraph_format.page_break_before = src.paragraph_format.page_break_before


def copy_runs(dst: Paragraph, src: Paragraph) -> None:
    for src_run in src.runs:
        run = dst.add_run(src_run.text)
        run.bold = src_run.bold
        run.italic = src_run.italic
        run.underline = src_run.underline
        run.font.superscript = src_run.font.superscript
        run.font.subscript = src_run.font.subscript
        run.font.size = src_run.font.size
        run.font.name = src_run.font.name
        east_asia = src_run._element.rPr.rFonts.get(qn("w:eastAsia")) if src_run._element.rPr is not None and src_run._element.rPr.rFonts is not None else None
        if east_asia and run._element.rPr is not None and run._element.rPr.rFonts is not None:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def insert_textual_paragraph_before(doc: Document, anchor, src: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor.addprevious(new_p)
    paragraph = Paragraph(new_p, doc._body)
    copy_paragraph_format(paragraph, src)
    copy_runs(paragraph, src)
    return paragraph


def replace_range_textually(
    dst_doc: Document,
    dst_start_idx: int,
    dst_end_idx: int,
    src_doc: Document,
    src_start_idx: int,
    src_end_idx: int,
) -> None:
    body = dst_doc._body._element
    dst_children = body_children(dst_doc)
    start = element_index(dst_doc, dst_start_idx)
    end = element_index(dst_doc, dst_end_idx)
    anchor = dst_children[end]

    for element in dst_children[start:end]:
        body.remove(element)

    for src_paragraph in src_doc.paragraphs[src_start_idx:src_end_idx]:
        insert_textual_paragraph_before(dst_doc, anchor, src_paragraph)


def replace_abstract_blocks(dst_doc: Document, src_doc: Document) -> None:
    zh_start_dst = find_text(dst_doc, "摘  要") + 1
    zh_end_dst = next(
        i
        for i in range(zh_start_dst, len(dst_doc.paragraphs))
        if dst_doc.paragraphs[i].text.strip().startswith("DESIGN AND IMPLEMENTATION")
    )
    zh_start_src = find_text(src_doc, "摘  要") + 1
    zh_end_src = next(
        i
        for i in range(zh_start_src, len(src_doc.paragraphs))
        if src_doc.paragraphs[i].text.strip().startswith("DESIGN AND IMPLEMENTATION")
    )
    replace_range_textually(dst_doc, zh_start_dst, zh_end_dst, src_doc, zh_start_src, zh_end_src)

    en_start_dst = find_text(dst_doc, "ABSTRACT") + 1
    en_end_dst = next(
        i
        for i in range(en_start_dst, len(dst_doc.paragraphs))
        if dst_doc.paragraphs[i].text.strip() == "目  录"
    )
    en_start_src = find_text(src_doc, "ABSTRACT") + 1
    en_end_src = next(
        i
        for i in range(en_start_src, len(src_doc.paragraphs))
        if src_doc.paragraphs[i].text.strip() == "目  录"
    )
    replace_range_textually(dst_doc, en_start_dst, en_end_dst, src_doc, en_start_src, en_end_src)


def replace_image_by_caption(doc: Document, caption_prefix: str, image_path: Path, width_inches: float) -> None:
    caption_idx = None
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith(caption_prefix):
            caption_idx = i
            break
    if caption_idx is None:
        raise ValueError(f"caption not found: {caption_prefix}")

    image_paragraph = None
    for j in range(caption_idx - 1, max(-1, caption_idx - 8), -1):
        if doc.paragraphs[j]._p.xpath(".//a:blip"):
            image_paragraph = doc.paragraphs[j]
            break
    if image_paragraph is None:
        raise ValueError(f"image paragraph not found before {caption_prefix}")

    clear_para(image_paragraph)
    image_paragraph.alignment = 1
    run = image_paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def ensure_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)


def main() -> None:
    shutil.copy2(V12, V14)
    src = Document(str(V11))
    dst = Document(str(V14))

    # Teacher comments: abstract replacement and keyword count.
    replace_abstract_blocks(dst, src)
    dst.save(str(V14))

    src = Document(str(V11))
    dst = Document(str(V14))

    # Teacher comments: Chapter 1 citations/cuts and Chapter 2 technology-only descriptions.
    replace_range_textually(
        dst,
        find_heading(dst, "1 绪论"),
        find_heading(dst, "3 系统分析"),
        src,
        find_heading(src, "1 绪论"),
        find_heading(src, "3 系统分析"),
    )

    # Teacher comments: generated engineering diagrams with small text.
    replacements = [
        ("图 3-1 用户角色权限关系图", "image1_roles_permissions.png", 5.65),
        ("图 3-2 系统业务流程图", "image2_business_flow.png", 5.70),
        ("图 3-3 告警工单闭环流程图", "image3_alert_loop.png", 5.65),
        ("图 3-4 系统用例图", "image4_use_case.png", 5.70),
        ("图 4-1 系统总体架构图", "image5_architecture.png", 5.70),
        ("图 4-2 系统部署结构图", "image6_deployment.png", 5.70),
        ("图 4-3 系统功能模块图", "image7_modules.png", 5.70),
        ("图 4-4 数据库 E-R 图", "image8_er.png", 5.55),
        ("图 4-5 系统核心类图", "image9_class.png", 5.55),
    ]
    for caption, filename, width in replacements:
        replace_image_by_caption(dst, caption, STABLE_DIAGRAMS / filename, width)

    ensure_update_fields(dst)
    dst.save(str(V14))
    print(V14)


if __name__ == "__main__":
    main()
