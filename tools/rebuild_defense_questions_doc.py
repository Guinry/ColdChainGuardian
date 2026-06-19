from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from append_defense_questions import SECTIONS


DOC_PATH = (
    Path.home()
    / "Desktop"
    / "毕设"
    / "答辩大语言模型相关问题整理.docx"
)


EXTRA_SECTIONS = [
    (
        "十、数据库与数据建模重点追问",
        [
            (
                "sensor_data 和 devices 为什么都保存温湿度，会不会数据冗余？",
                "这是有意的冗余设计。sensor_data 保存每一次采集的历史明细，用于趋势图、追溯和统计；devices 保存 latest_temp、latest_humi、latest_data_time，用于设备列表、首页和实时监测快速读取。这样可以减少每次都从历史表聚合最新记录的开销。关键是上报时要在同一个业务流程中同时写历史表并更新设备最新状态，保证一致性。",
            ),
            (
                "为什么使用外键？外键带来了什么问题？",
                "外键可以保证 alerts.device_id、work_orders.alert_id 等字段引用的数据真实存在，避免出现告警指向不存在设备的脏数据。它的代价是导入演示数据、删除设备或重建数据库时必须注意顺序和 ID 映射，否则会出现外键约束失败。答辩时可以说明：外键是为了数据可靠性，线上问题主要来自演示数据重建后旧 ID 未同步，修复思路是按业务编码重建映射并避免整行写回旧外键值。",
            ),
            (
                "如果 sensor_data 数据量越来越大怎么办？",
                "可以从四个方向优化：第一，对 device_id 和 data_time 建联合索引；第二，按月份或按设备分表，减少单表数据量；第三，对首页图表做按分钟、小时的预聚合；第四，长期可引入时序数据库，例如 InfluxDB 或 TimescaleDB。毕业设计阶段数据规模较小，MySQL 加索引已经足够。",
            ),
            (
                "MySQL 5.7 和 MySQL 8 兼容问题体现在哪里？",
                "部署服务器使用 MySQL 5.7，部分 MySQL 8 导出的 SQL 语法和排序规则不能直接导入，例如 utf8mb4_0900_ai_ci 需要替换为 utf8mb4_unicode_ci。答辩时可以说：部署前对 SQL 做了兼容处理，保证表结构和演示数据能在服务器环境稳定运行。",
            ),
            (
                "为什么库区表使用 parent_id，而不是单独建多张区域表？",
                "parent_id 可以表达任意层级的树形结构，例如仓库、子库区、货架等，不需要每增加一层就修改表结构。前端可以按 parent_id 组装树，后端也能按树结构统计设备和告警。缺点是查询完整树需要额外组装，必要时可以增加 area_path 提高查询效率。",
            ),
            (
                "如何保证告警转工单不会重复创建？",
                "可以在业务层判断该告警是否已经有关联工单，或者在 work_orders.alert_id 上增加唯一约束。当前演示流程中，转工单时会读取告警状态，创建工单后更新告警状态为 HANDLING。更严格的生产方案应使用事务和唯一约束共同保证幂等性。",
            ),
            (
                "为什么有些接口返回 200 但 success=false？",
                "这是统一响应格式的设计。HTTP 200 表示请求成功到达后端并被业务处理，success=false 表示业务逻辑失败，例如找不到员工档案或转工单失败。这样前端能统一读取 message 给用户提示。生产中也可以把部分业务错误映射为 4xx 状态码。",
            ),
        ],
    ),
    (
        "十一、后端、安全与权限追问",
        [
            (
                "JWT 的作用是什么？为什么小程序会出现 403？",
                "JWT 用于登录后的身份认证。用户登录成功后，后端签发 token，前端或小程序后续请求在 Authorization 请求头携带 token。403 通常表示 token 缺失、过期、角色权限不足，或小程序登录接口返回的用户角色没有访问对应接口的权限。排查时先看登录返回 token，再看请求头是否携带 Authorization，最后看 SecurityConfig 中对应路径允许的角色。",
            ),
            (
                "密码为什么不能明文保存？",
                "明文密码一旦数据库泄露会直接造成账号风险。正确做法是使用 BCrypt 等不可逆哈希保存密码摘要，登录时用输入密码和哈希值进行匹配。演示账号统一为 123456 只是为了答辩方便，不代表生产环境应该使用弱密码。",
            ),
            (
                "如何防止 SQL 注入？",
                "项目使用 MyBatis/MyBatis-Plus 的参数绑定、LambdaQueryWrapper 等方式构造查询，避免把用户输入直接拼接到 SQL 字符串中。同时对分页、筛选、状态等参数做枚举或类型校验，可以进一步降低 SQL 注入风险。",
            ),
            (
                "为什么 AI Key、数据库密码不能写在前端？",
                "前端代码和小程序包都可能被用户看到，如果把 API Key 或数据库密码写进去，就等于公开泄露。正确做法是把敏感配置放在服务器环境变量或 root-only 环境文件中，由后端读取，前端只调用自己的后端接口。",
            ),
            (
                "后端如何处理异常？",
                "业务异常会被转换成统一的 ApiResponse，包括 success、code、message、data。对于数据库异常、外部 AI 服务异常等，后端记录日志并返回明确提示。这样前端、小程序和调试人员都能用一致方式判断问题。",
            ),
            (
                "Spring Boot 分层设计有什么好处？",
                "Controller 负责接口入口，Service 负责业务逻辑，Mapper/Repository 负责数据库访问，DTO 负责接口数据传输。分层后代码职责清晰，便于维护和排查。例如 AI 助手相关逻辑集中在 AIAssistantService 和 AiModelClient，IoT 上报逻辑集中在 IotTelemetryService。",
            ),
            (
                "Redis 在项目里可以承担什么作用？",
                "Redis 适合做缓存、短期状态、验证码、会话或热点数据加速。当前毕业设计核心数据仍以 MySQL 为准，Redis 不是强依赖业务库。答辩时可以说：后期如果实时监测设备数量扩大，可以把设备最新状态和首页统计放入 Redis，降低数据库压力。",
            ),
        ],
    ),
    (
        "十二、部署、域名与服务器追问",
        [
            (
                "线上访问链路是怎样的？",
                "用户访问 https://coldchain.guinry.cn，DNS A 记录解析到阿里云服务器 47.104.195.63。Nginx 接收 80/443 请求，静态页面直接从前端 dist 目录返回，/api 请求反向代理到本机 127.0.0.1:18080 的 Spring Boot 服务，数据库 MySQL 和 Redis 不直接暴露公网。",
            ),
            (
                "为什么后端只监听 127.0.0.1:18080？",
                "这样可以避免后端端口直接暴露公网，外部只能通过 Nginx 访问。Nginx 统一处理 HTTPS、静态资源、反向代理和日志，安全边界更清晰。公网最终只需要保留 22、80、443。",
            ),
            (
                "为什么必须配置 HTTPS？",
                "HTTPS 可以加密登录信息、JWT token 和业务数据，防止传输过程中被窃听或篡改。微信小程序正式请求服务器也要求使用 HTTPS 合法域名，所以 HTTPS 是部署验收的重要条件。",
            ),
            (
                "2 核 2G 的服务器够用吗？",
                "对于毕业设计演示、小规模 IoT 设备上报和少量用户访问是够用的。关键是控制 JVM 内存，例如设置 Xms128m、Xmx650m，前端静态资源交给 Nginx，数据库和后端在本机内网访问。若后期设备量、用户量或 AI 请求量增加，可以升级内存或拆分服务。",
            ),
            (
                "为什么清理宝塔、FTP、MQTT、frps、tinyproxy 等服务？",
                "这些服务如果当前项目不用，却在公网监听，会增加攻击面和维护复杂度。清理后服务器更干净，排查问题也更简单。项目实际需要的是 Nginx、Spring Boot、MySQL，以及必要时的 Redis，其他服务应关闭或只内网使用。",
            ),
            (
                "Nginx 出现 502 应该怎么排查？",
                "先检查 Spring Boot 是否启动，执行 systemctl status coldchain-guardian；再 curl http://127.0.0.1:18080/actuator/health 看本机后端是否可用；然后检查 Nginx 反代配置和错误日志。502 通常是后端未启动、端口不一致或 Nginx 代理地址写错。",
            ),
            (
                "为什么前端生产环境 API 要使用 /api 相对路径？",
                "使用 /api 可以让前端和后端保持同源访问，由 Nginx 负责转发，避免浏览器跨域问题，也避免把 localhost 或内网地址打包进生产前端。部署后访问 coldchain.guinry.cn 时，所有接口都自然走同一个域名。",
            ),
        ],
    ),
    (
        "十三、小程序、前端与演示追问",
        [
            (
                "小程序的请求地址改成线上了吗？",
                "已经改为 https://coldchain.guinry.cn。小程序 utils/request.js 中默认 BASE_URL 是线上域名，并且会把以前保存的 localhost、127.0.0.1、192.168.x.x 这类本地地址自动忽略，避免演示时还请求本机。",
            ),
            (
                "微信开发者工具里的 Skyline、getSystemInfo 警告影响演示吗？",
                "这些主要是开发者工具的兼容性提示和 API 废弃提醒，不等于业务错误。演示时重点看接口请求是否成功、页面数据是否正常。如果要进一步优化，可以把 getSystemInfoSync 替换为新的 wx.getDeviceInfo、wx.getWindowInfo 等 API。",
            ),
            (
                "小程序登录提示未找到手机号对应员工档案怎么办？",
                "说明后端能收到登录请求，但数据库里没有与该手机号绑定的员工或用户档案。解决方式是在 users 或员工相关表里补齐测试账号和手机号，或者使用已存在的演示账号登录。答辩时可以说明这是业务数据校验，不是接口连不通。",
            ),
            (
                "前端为什么使用 Vue/Vite？",
                "Vue 适合构建管理后台的组件化页面，Vite 启动和构建速度快，开发体验好。系统包含仪表盘、表格、树结构、图表、弹窗和 AI 对话等多个组件，使用 Vue 能让页面结构更清晰。",
            ),
            (
                "AI 助手为什么使用 SSE？",
                "AI 回复通常是逐步生成的，SSE 可以让前端边生成边显示，用户不用等待完整回答结束。相比轮询，SSE 更简单；相比 WebSocket，SSE 对单向文本流式输出更轻量。",
            ),
            (
                "答辩演示时页面数据突然不显示怎么办？",
                "按三步处理：先刷新页面并确认已登录；再打开浏览器网络面板看接口状态码；如果是 401/403，重新登录；如果是 502，检查后端服务；如果接口 200 但页面无数据，说明是数据为空或字段映射问题，可以切到数据库或接口返回说明系统链路正常。",
            ),
        ],
    ),
    (
        "十四、IoT 硬件与传感器追问",
        [
            (
                "为什么最终选择 OLED，而不是裸 16 脚 LCD1602？",
                "裸 16 脚 LCD1602 接线多，需要对比度电位器、背光供电和较多 GPIO，对于只有杜邦线、没有面包板的 ESP32 演示环境不方便。0.96 寸 SSD1306 OLED 只有 GND、VCC、SCL、SDA 四根线，可以和 SHT31 共用 I2C 总线，更适合当前硬件条件。",
            ),
            (
                "SHT31 和 OLED 为什么可以共用 SDA/SCL？",
                "它们都是 I2C 设备，I2C 总线允许多个设备共用 SDA 和 SCL，只要设备地址不同即可。SHT31 常用地址是 0x44，SSD1306 OLED 常用地址是 0x3C，因此可以共用 GPIO21 和 GPIO22。",
            ),
            (
                "为什么 SHT31 和 OLED 都建议接 3.3V？",
                "ESP32 的 GPIO 不是 5V 容忍的。很多 I2C 模块会把 SDA/SCL 上拉到 VCC，如果模块接 5V，可能把 I2C 线拉到 5V，损坏 ESP32 或导致不稳定。因此在没有电平转换模块时，SHT31 和 OLED 都接 3.3V 更安全。",
            ),
            (
                "OLED 黑屏怎么排查？",
                "先检查 VCC/GND 是否接反，再检查 SDA/SCL 是否接到 GPIO21/GPIO22；然后看串口是否打印 SSD1306 OLED ready；如果地址不对，尝试 0x3C 和 0x3D；如果仍黑屏，可能是供电、排针焊接或屏幕本身问题。代码里已经支持初始化失败提示和地址 fallback。",
            ),
            (
                "SHT31 read failed 怎么排查？",
                "先检查 VCC、GND、SDA、SCL 是否接牢，再确认 SHT31 地址是 0x44 还是 0x45；如果 OLED 和 SHT31 共线，检查分线是否接触不良；还要避免把 SHT31 接 5V 导致 I2C 电平不匹配。串口中 SHT31 ready 说明硬件识别成功，read failed 则可能是线松或瞬时通信失败。",
            ),
            (
                "为什么使用 HTTP 上报，不使用 MQTT？",
                "当前项目设备数量少、上报频率低，HTTP POST 到 /api/iot/telemetry 更直接，便于和 Spring Boot REST API 对接，也方便答辩演示和调试。MQTT 更适合大量设备、低延迟消息推送和双向控制，后期可以扩展。",
            ),
            (
                "设备离线时还能显示温湿度吗？",
                "可以。ESP32 本地读取 SHT31 后会在 OLED 上显示最新温湿度，网络失败只影响上传到服务器，不影响本地采集显示。这样答辩现场即使网络波动，也能证明硬件采集链路正常。",
            ),
        ],
    ),
    (
        "十五、答辩现场故障排查快答",
        [
            (
                "AI 页面显示“外部模型服务暂时不可用”怎么解释？",
                "说明系统已经从数据库读取了运行数据，也能做本地规则分析，只是外部大模型接口暂时不可用。可以强调 AI 模块有降级机制，不影响核心的温湿度监测、告警、工单和数据管理。",
            ),
            (
                "告警转工单失败，外键约束报错怎么讲？",
                "这是数据一致性问题，不是流程设计问题。因为演示数据库重建后部分告警的 device_id 指向旧设备 ID，触发了外键保护。解决方案是修复历史演示数据的 device_id 映射，并在代码中只更新告警状态字段，避免把旧外键整行写回。",
            ),
            (
                "库区管理或实时监测看不到子库区怎么讲？",
                "这是树形数据 parent_id 或 area_path 与当前父节点不一致造成的。解决思路是按 area_code 修正父子关系，并让后端构建树时对孤立节点做兜底展示，避免页面看不到数据。",
            ),
            (
                "小程序请求 403 怎么讲？",
                "403 表示请求到了服务器，但权限校验未通过。可能原因是 token 未携带、token 过期、角色不匹配，或登录用户没有访问该接口的权限。排查重点是登录返回、请求头 Authorization 和后端 SecurityConfig。",
            ),
            (
                "ESP32 上报失败怎么讲？",
                "先看 OLED 和串口是否仍显示温湿度。如果本地数据正常，只是 upload failed，说明传感器链路正常，问题在 Wi-Fi、域名、HTTPS、后端接口或服务器网络。可以现场用串口日志展示 POST 地址和返回状态。",
            ),
            (
                "线上页面打不开怎么讲？",
                "先检查 DNS 是否解析到 47.104.195.63，再检查 Nginx 是否运行、443 证书是否有效、后端是否启动。答辩时可以准备本地截图或录屏作为兜底，但优先演示线上地址。",
            ),
            (
                "老师问项目不足怎么回答？",
                "可以主动承认三个方向：第一，当前设备规模较小，后期需要 MQTT、离线缓存和批量设备管理；第二，历史数据量增大后需要分表或时序数据库；第三，AI 模块目前以结构化数据库上下文为主，后续可以加入向量知识库、更多冷链规范文档和更严格的结果校验。",
            ),
        ],
    ),
]


def set_font(run, size: float = 10.5, bold: bool = False) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def set_heading_font(paragraph, size: float = 15, bold: bool = True) -> None:
    for run in paragraph.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        run.font.size = Pt(size)
        run.bold = bold


def add_question(document: Document, question: str, answer: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    q = p.add_run("问：" + question)
    set_font(q, bold=True)

    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    a = p.add_run("答：" + answer)
    set_font(a)


def add_section(document: Document, heading: str, qas: list[tuple[str, str]]) -> None:
    p = document.add_heading(heading, level=1)
    set_heading_font(p, size=14, bold=True)
    for question, answer in qas:
        add_question(document, question, answer)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)


def build_document() -> Document:
    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ColdChain Guardian 毕业答辩可能问题与参考回答")
    set_heading_font(title, size=18, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("覆盖大语言模型、数据库、后端、小程序、IoT、部署、安全与现场排查")
    set_font(run, size=10.5)

    document.add_paragraph()
    intro = document.add_paragraph()
    intro.paragraph_format.first_line_indent = Pt(21)
    intro.paragraph_format.line_spacing = 1.15
    run = intro.add_run(
        "使用建议：答辩时不需要逐字背诵，先用一句话回答结论，再结合本项目的真实实现举例。"
        "如果老师追问细节，优先从“为什么这样设计、如何实现、出现问题如何排查、后续如何优化”四个角度展开。"
    )
    set_font(run)

    for heading, qas in list(SECTIONS) + EXTRA_SECTIONS:
        add_section(document, heading, qas)

    return document


def main() -> None:
    if DOC_PATH.exists():
        backup = DOC_PATH.with_name(
            DOC_PATH.stem + ".before-rebuild-" + datetime.now().strftime("%Y%m%d-%H%M%S") + ".docx"
        )
        shutil.copy2(DOC_PATH, backup)
    else:
        backup = None

    document = build_document()
    document.save(DOC_PATH)

    with ZipFile(DOC_PATH) as archive:
        xml = archive.read("word/document.xml").decode("utf-8", errors="replace")

    chinese_chars = sum("\u4e00" <= char <= "\u9fff" for char in xml)
    if chinese_chars < 1000:
        raise RuntimeError("Document verification failed: Chinese content was not written correctly.")

    print(f"updated={DOC_PATH}")
    if backup:
        print(f"backup={backup}")
    print(f"chinese_chars={chinese_chars}")
    print(f"question_count={xml.count('问：')}")


if __name__ == "__main__":
    main()
