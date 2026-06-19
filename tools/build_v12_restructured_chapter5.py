from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SRC = Path(r"C:\Users\23869\Desktop\毕设\草稿\郭鑫瑞毕业论文v9.docx")
DST = Path(r"C:\Users\23869\Desktop\毕设\草稿\郭鑫瑞毕业论文v12.docx")
ROOT = Path(r"C:\Users\23869\Desktop\ColdChainGuardian")
ASSETS = ROOT / "output" / "doc" / "qa_v12"
DIAGRAMS = ASSETS / "module_diagrams"
SNAPS = ASSETS / "selected_screenshots"


def set_border(element, edge: str, val: str = "single", size: str = "8") -> None:
    tag = qn(f"w:{edge}")
    child = element.find(tag)
    if child is None:
        child = OxmlElement(f"w:{edge}")
        element.append(child)
    child.set(qn("w:val"), val)
    child.set(qn("w:sz"), size)
    child.set(qn("w:space"), "0")
    child.set(qn("w:color"), "000000")


def clear_border(element, edge: str) -> None:
    tag = qn(f"w:{edge}")
    child = element.find(tag)
    if child is None:
        child = OxmlElement(f"w:{edge}")
        element.append(child)
    child.set(qn("w:val"), "nil")


def apply_three_line_table(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("left", "right", "insideV"):
        clear_border(borders, edge)
    for edge in ("top", "bottom"):
        set_border(borders, edge, size="12")
    clear_border(borders, "insideH")
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.first_child_found_in("w:tcBorders")
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        set_border(tc_borders, "bottom", size="8")


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


class ChapterBuilder:
    def __init__(self, doc: Document, anchor):
        self.doc = doc
        self.anchor = anchor

    def _move_before_anchor(self, element) -> None:
        self.anchor.addprevious(element)

    def paragraph(self, text: str = "", style: str | None = None, align: int | None = None):
        p = self.doc.add_paragraph()
        if style:
            p.style = style
        if text:
            p.add_run(text)
        if align is not None:
            p.alignment = align
        if style == "Heading 1":
            p.paragraph_format.page_break_before = True
        self._move_before_anchor(p._p)
        return p

    def body(self, text: str):
        p = self.paragraph(text)
        p.style = "Normal"
        return p

    def heading(self, level: int, text: str):
        return self.paragraph(text, f"Heading {level}")

    def image(self, path: Path, width: float = 5.7):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
        self._move_before_anchor(p._p)
        return p

    def caption(self, cn: str, en: str):
        p1 = self.paragraph(cn, align=WD_ALIGN_PARAGRAPH.CENTER)
        p2 = self.paragraph(en, align=WD_ALIGN_PARAGRAPH.CENTER)
        for p in (p1, p2):
            for run in p.runs:
                run.font.name = "Times New Roman" if p is p2 else "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(10.5)
        return p1, p2

    def table(self, headers: list[str], rows: list[list[str]], cn: str, en: str):
        self.caption(cn, en)
        table = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, header in enumerate(headers):
            set_cell_text(table.cell(0, i), header, bold=True)
        for r, row in enumerate(rows, start=1):
            for c, value in enumerate(row):
                set_cell_text(table.cell(r, c), value)
        apply_three_line_table(table)
        self._move_before_anchor(table._tbl)
        return table


def find_paragraph(doc: Document, prefix: str):
    for p in doc.paragraphs:
        if p.style and p.style.name == "Heading 1" and (p.text or "").strip().startswith(prefix):
            return p
    raise ValueError(f"Paragraph not found: {prefix}")


def remove_chapter5(doc: Document):
    h5 = find_paragraph(doc, "5 系统详细设计与实现")
    h6 = find_paragraph(doc, "6 系统测试")
    body = doc._body._element
    children = list(body)
    start = children.index(h5._p)
    end = children.index(h6._p)
    for element in children[start:end]:
        body.remove(element)
    return h6._p


def build_chapter5(builder: ChapterBuilder) -> None:
    b = builder
    b.heading(1, "5 系统详细设计与实现")
    b.body("系统详细设计与实现章节按照核心业务模块展开。每个模块均将详细设计、关键流程或时序图、源代码实现说明和运行界面效果放在同一小节中，避免设计说明与实现截图相互分离。页面实现效果并非按菜单完整堆叠，而是选择能体现主要功能闭环的代表性界面。")

    b.heading(2, "5.1 系统实现环境与工程组织")
    b.heading(3, "5.1.1 开发与运行环境")
    b.body("系统采用前后端分离和移动端协同的实现方式。后端基于 Spring Boot 提供业务接口、鉴权、实时监测、告警工单和智能助手服务；Web 管理端基于 Vue 3、Element Plus 和 ECharts 实现管理页面；微信小程序面向现场员工提供移动工单和告警处理入口。测试阶段使用 MySQL 存储业务数据，Redis 用于缓存和状态辅助，AI 助手通过兼容式大模型接口接入 deepseek-v4-pro。")
    b.table(
        ["类别", "实现环境", "说明"],
        [
            ["后端服务", "Spring Boot、MyBatis-Plus、JWT", "提供认证、业务接口、异常处理和日志记录"],
            ["Web 管理端", "Vue 3、Element Plus、ECharts", "实现后台管理、图表展示和复杂表格交互"],
            ["小程序端", "微信小程序原生框架", "实现现场员工移动处理告警和工单"],
            ["数据与缓存", "MySQL、Redis", "存储库区、设备、监测数据、告警、工单和 AI 会话"],
            ["智能助手", "OpenAI-compatible API、deepseek-v4-pro", "按业务上下文组织提示并返回分析建议"],
        ],
        "表 5-1 系统开发与运行环境",
        "Table 5-1 Development and runtime environment of the system",
    )

    b.heading(3, "5.1.2 工程目录与源代码组织")
    b.body("工程目录按照后端、Web 管理端和小程序端划分。后端 ccg-app 模块承担 Controller、Service、Scheduler 和安全配置职责，ccg-contract 模块保存 DTO、VO 与枚举定义，ccg-common 模块提供统一响应和异常结构。Web 端以 views、api、components、router 和 store 组织页面、接口封装和状态控制。小程序端以 pages/workbench、pages/alert、pages/workorder 和 pages/profile 组织移动端页面。")
    b.table(
        ["模块", "后端源代码", "前端或小程序源代码"],
        [
            ["Dashboard 与认证", "AuthController、DashboardController、DashboardService", "LoginPage.vue、Dashboard.vue、router/index.js"],
            ["库区与设备", "AreaController、DeviceController、AreaService、DeviceService", "WarehouseAreaManage.vue、DeviceManagementView.vue、api/area.js、api/device.js"],
            ["实时监测与阈值", "MonitorController、TelemetryService、AlertScheduler", "MonitorView.vue、ThresholdSettingsView.vue、api/monitor.js"],
            ["告警与工单", "AlertController、WorkOrderController、AlertService、WorkOrderService", "AlertCenterView.vue、WorkOrderCenterView.vue、pages/workorder"],
            ["AI 智能助手", "AIAssistantService、AiModelClient、AlertAnalysisService", "AIAssistantView.vue、components/DataTableCard.vue"],
            ["小程序端", "WxAuthService、WorkOrderService、AlertService", "pages/workbench、pages/alert、pages/workorder、pages/profile"],
        ],
        "表 5-2 核心模块源代码对应关系",
        "Table 5-2 Mapping between core modules and source code",
    )

    b.heading(2, "5.2 Dashboard 与登录认证模块详细设计与实现")
    b.heading(3, "5.2.1 模块详细设计")
    b.body("Dashboard 是 Web 管理端的首页，也是用户进入系统后获取运行态势的第一入口。登录认证流程先完成账号密码校验和角色权限加载，后端生成 Token 后返回前端，前端路由守卫根据登录状态和角色信息决定是否允许进入对应页面。Dashboard 页面加载时会请求设备统计、告警统计、趋势数据和待处理事项，并将结果组织为指标卡片、趋势图和列表提醒。")
    b.image(DIAGRAMS / "fig5_01_auth_dashboard_sequence.png", 5.7)
    b.caption("图 5-1 登录认证与 Dashboard 加载时序图", "Figure 5-1 Sequence diagram of login authentication and Dashboard loading")
    b.heading(3, "5.2.2 源代码实现与运行效果")
    b.body("后端认证逻辑主要由 AuthController 和 AuthService 完成，首页统计数据由 DashboardController 调用 DashboardService 汇总。前端 LoginPage.vue 负责登录表单提交，router/index.js 完成路由守卫，Dashboard.vue 负责展示指标卡片、告警趋势、库区状态和待处理提醒。该页面实现效果如图 5-2 所示。")
    b.image(SNAPS / "dashboard.png", 5.75)
    b.caption("图 5-2 Dashboard 页面实现效果", "Figure 5-2 Implementation effect of Dashboard page")

    b.heading(2, "5.3 库区与设备管理模块详细设计与实现")
    b.heading(3, "5.3.1 模块详细设计")
    b.body("库区与设备管理模块负责维护冷链仓储空间结构和监测终端基础信息。库区管理采用树形结构表达父子层级，页面选择左侧库区节点后加载基础信息、子库区和关联设备；设备管理围绕设备编号、所属库区、在线状态、传感类型和启停状态展开。新增或编辑设备时，后端需要校验所属库区、设备编号唯一性和状态合法性，保存后刷新列表和关联库区数据。")
    b.image(DIAGRAMS / "fig5_03_area_device_flow.png", 5.75)
    b.caption("图 5-3 库区与设备管理模块流程图", "Figure 5-3 Flow diagram of warehouse area and device management module")
    b.heading(3, "5.3.2 库区管理实现效果")
    b.body("库区管理页面由 WarehouseAreaManage.vue 及其子组件组成，左侧显示库区结构，右侧显示基础信息、子库区、设备数量和维护操作。该页面将树形浏览与详情维护放在同一工作区，便于管理员按照库区层级逐级维护。")
    b.image(SNAPS / "area_management.png", 5.75)
    b.caption("图 5-4 库区管理页面实现效果", "Figure 5-4 Implementation effect of warehouse area management page")
    b.heading(3, "5.3.3 设备管理实现效果")
    b.body("设备管理页面由 DeviceManagementView.vue 和 api/device.js 共同实现。页面支持按状态、库区和关键字筛选设备，表格中展示在线状态、所属库区、最近上报时间和启停状态，并提供新增、编辑、启用、停用和查看数据等操作。")
    b.image(SNAPS / "device_management.png", 5.75)
    b.caption("图 5-5 设备管理页面实现效果", "Figure 5-5 Implementation effect of device management page")

    b.heading(2, "5.4 实时监测与阈值规则模块详细设计与实现")
    b.heading(3, "5.4.1 模块详细设计")
    b.body("实时监测模块以设备采集数据为输入，后端接收温度、湿度、电量、在线状态等监测值后写入数据库，并与阈值规则进行匹配。若数据超出规则范围，系统将更新告警状态并通过页面刷新或 WebSocket 推送提醒前端。阈值规则维护页面用于配置温湿度范围、设备离线时间和异常级别，为监测数据判断提供依据。")
    b.image(DIAGRAMS / "fig5_06_monitor_threshold_flow.png", 5.75)
    b.caption("图 5-6 实时监测与阈值规则模块流程图", "Figure 5-6 Flow diagram of real-time monitoring and threshold rule module")
    b.heading(3, "5.4.2 实时监测实现效果")
    b.body("实时监测页面由 MonitorView.vue、MonitorKpiCards.vue、RealtimeDeviceTable.vue 和 DeviceRealtimeDrawer.vue 组成，页面同时展示关键指标、设备状态表格和趋势信息。页面实现效果如图 5-7 所示。")
    b.image(SNAPS / "realtime_monitor.png", 5.75)
    b.caption("图 5-7 实时监测页面实现效果", "Figure 5-7 Implementation effect of real-time monitoring page")
    b.heading(3, "5.4.3 阈值规则实现效果")
    b.body("阈值规则页面由 ThresholdSettingsView.vue 实现，管理员可以维护不同监测指标的上下限、告警级别和启用状态。规则变化后，后端在监测数据处理和告警生成阶段统一读取规则，保证页面配置能够直接影响业务判断。")
    b.image(SNAPS / "threshold_rules.png", 5.75)
    b.caption("图 5-8 阈值规则页面实现效果", "Figure 5-8 Implementation effect of threshold rule page")

    b.heading(2, "5.5 告警中心与工单闭环模块详细设计与实现")
    b.heading(3, "5.5.1 模块详细设计")
    b.body("告警中心与工单闭环模块是系统业务闭环的核心。异常数据触发告警后，仓储管理员在告警中心进行确认、忽略、关闭或转工单操作；需要现场处理的告警会生成工单并分配给现场员工。员工在小程序端接收任务、查看详情、处理现场问题并提交反馈，管理员验收通过后关闭归档；若验收不通过，工单退回继续处理。")
    b.image(DIAGRAMS / "fig5_09_alert_order_flow.png", 5.75)
    b.caption("图 5-9 告警与工单闭环模块流程图", "Figure 5-9 Flow diagram of alert and work order closed-loop module")
    b.heading(3, "5.5.2 告警中心实现效果")
    b.body("告警中心页面由 AlertCenterView.vue 和 AlertTriageDrawer.vue 实现，页面提供告警等级、处理状态、库区设备和时间条件筛选，支持告警确认、关闭、转工单和查看处置记录。")
    b.image(SNAPS / "alert_center.png", 5.75)
    b.caption("图 5-10 告警中心页面实现效果", "Figure 5-10 Implementation effect of alert center page")
    b.heading(3, "5.5.3 工单中心实现效果")
    b.body("工单中心页面由 WorkOrderCenterView.vue、CreateWorkOrderModal.vue 和 WorkOrderDrawer.vue 组成，围绕工单状态流转展示待接收、处理中、待验收和已完成任务，并提供派发、查看、验收和关闭等管理操作。")
    b.image(SNAPS / "work_order_center.png", 5.75)
    b.caption("图 5-11 工单中心页面实现效果", "Figure 5-11 Implementation effect of work order center page")

    b.heading(2, "5.6 数据分析与系统管理模块详细设计与实现")
    b.heading(3, "5.6.1 模块详细设计")
    b.body("数据分析模块面向管理人员提供监测趋势、告警数量、设备状态和库区风险统计。系统管理模块负责用户、角色、权限和审计日志维护。两个模块虽然业务目标不同，但都依赖后端对数据库进行查询聚合，并通过前端表格、指标卡片和图表展示结果。权限校验贯穿查询与管理操作全过程，审计日志记录关键变更。")
    b.image(DIAGRAMS / "fig5_12_analysis_system_flow.png", 5.75)
    b.caption("图 5-12 数据分析与系统管理模块流程图", "Figure 5-12 Flow diagram of data analysis and system management module")
    b.heading(3, "5.6.2 数据分析实现效果")
    b.body("数据分析页面由 TrendAnalysisView.vue 及其图表组件组成，页面通过指标卡片、趋势曲线和明细表格展示冷链运行情况，为管理人员判断异常高发时段和风险库区提供依据。")
    b.image(SNAPS / "data_analysis.png", 5.75)
    b.caption("图 5-13 数据分析页面实现效果", "Figure 5-13 Implementation effect of data analysis page")
    b.heading(3, "5.6.3 系统管理实现效果")
    b.body("系统管理页面重点体现用户与权限维护。PermissionManagementView.vue 负责角色、菜单和操作权限展示，后端根据角色信息控制接口访问范围，并通过审计日志记录关键操作。")
    b.image(SNAPS / "system_permission.png", 5.75)
    b.caption("图 5-14 系统管理页面实现效果", "Figure 5-14 Implementation effect of system management page")

    b.heading(2, "5.7 AI 智能助手模块详细设计与实现")
    b.heading(3, "5.7.1 模块详细设计")
    b.body("AI 智能助手模块用于将业务数据库中的库区、设备、监测数据、告警和工单信息整理为上下文，再通过兼容式大模型接口生成风险解释和处置建议。前端提交自然语言问题后，后端先识别问题意图，再查询相关业务数据并构造提示词，模型返回结果后由前端以 Markdown 结构化渲染。")
    b.image(DIAGRAMS / "fig5_15_ai_assistant_sequence.png", 5.75)
    b.caption("图 5-15 AI 智能助手调用时序图", "Figure 5-15 Sequence diagram of AI assistant invocation")
    b.heading(3, "5.7.2 源代码实现与运行效果")
    b.body("后端 AIAssistantService 负责会话保存、上下文构造和模型调用，AiModelClient 负责兼容式接口请求，AlertAnalysisService 负责从告警和监测数据中提取分析素材。前端 AIAssistantView.vue 负责会话列表、消息输入、加载状态和结果渲染。页面实现效果如图 5-16 所示。")
    b.image(SNAPS / "ai_assistant.png", 5.75)
    b.caption("图 5-16 AI 智能助手页面实现效果", "Figure 5-16 Implementation effect of AI assistant page")

    b.heading(2, "5.8 微信小程序端详细设计与实现")
    b.heading(3, "5.8.1 模块详细设计")
    b.body("微信小程序端面向现场员工，主要承担告警查看、工单接收、现场处置、反馈提交和个人中心功能。小程序登录后进入工作台，用户可查看紧急告警和待处理工单；进入工单页面后按状态筛选任务，接收后提交处理反馈，待管理员验收后进入完成状态。")
    b.image(DIAGRAMS / "fig5_18_mini_program_flow.png", 5.75)
    b.caption("图 5-17 小程序端任务处理流程图", "Figure 5-17 Flow diagram of mini program task handling")
    b.heading(3, "5.8.2 工作台、告警与工单实现效果")
    b.body("小程序工作台、告警大厅和工单处理页面分别对应 pages/workbench、pages/alert 和 pages/workorder 目录。工作台展示紧急告警、设备状态和快捷入口，告警大厅区分未处理、处理中和已恢复状态，工单页面区分待接收与处理中任务。")
    b.image(SNAPS / "mini_workbench_alert_order.png", 5.75)
    b.caption("图 5-18 小程序工作台、告警与工单页面运行效果", "Figure 5-18 Running effect of mini program workbench, alert and work order pages")
    b.heading(3, "5.8.3 验收、已完成与个人中心实现效果")
    b.body("小程序端还提供待验收、已完成和个人中心页面。待验收页面用于员工查看已提交但尚未验收的任务，已完成页面展示闭环结果，个人中心展示工单统计、连接状态和消息提醒开关。")
    b.image(SNAPS / "mini_verify_profile.png", 5.75)
    b.caption("图 5-19 小程序验收、已完成与个人中心页面运行效果", "Figure 5-19 Running effect of mini program verification, completed and profile pages")

    b.heading(2, "5.9 本章小结")
    b.body("本章围绕系统核心功能完成了详细设计与实现说明。与概要设计阶段不同，本章将每个重要模块的处理流程、源代码实现位置和页面效果放在同一模块小节中进行说明，使需求、设计、实现和运行效果形成对应关系。通过 Dashboard、库区设备、实时监测、告警工单、数据分析、系统管理、AI 智能助手和微信小程序端的实现，可以支撑冷链仓储安全管理从异常发现到处置闭环的主要业务过程。")


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    DST.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))
    anchor = remove_chapter5(doc)
    builder = ChapterBuilder(doc, anchor)
    build_chapter5(builder)
    doc.save(str(DST))
    print(DST)


if __name__ == "__main__":
    main()
