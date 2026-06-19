from __future__ import annotations

import base64
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def decode_arg(value: str) -> Path:
    return Path(base64.b64decode(value).decode("utf-8"))


def clear_para(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def set_text(paragraph, text: str) -> None:
    clear_para(paragraph)
    if text:
        paragraph.add_run(text)


def set_text_with_cites(paragraph, parts) -> None:
    clear_para(paragraph)
    for item in parts:
        if isinstance(item, tuple):
            run = paragraph.add_run(item[0])
            if item[1] == "sup":
                run.font.superscript = True
        else:
            paragraph.add_run(item)


def main() -> int:
    src = decode_arg(sys.argv[1])
    dst = decode_arg(sys.argv[2])
    shutil.copy2(src, dst)
    doc = Document(str(dst))

    abstract = (
        "冷链仓储是保障生鲜食品、医药疫苗及其他温控敏感物资质量安全的重要环节，在现代物流体系中具有较高的时效性与安全性要求。随着冷链物流规模不断扩大，传统冷链仓储管理方式在实时监测、异常响应、工单闭环以及数据智能分析等方面逐渐暴露出不足，难以满足复杂业务场景下的安全管理需求。同时，现有系统普遍存在数据展示与业务处置脱节、移动端协同能力不足以及智能分析能力有限等问题。针对上述问题，本文以冷链仓储安全管理为研究对象，设计并实现一套基于大语言模型的冷链仓储安全管理系统，以提升冷链仓储异常发现效率、业务协同能力及智能化分析水平。\n\n"
        "本文采用前后端分离架构进行系统设计与实现。后端基于 Spring Boot、MySQL、Redis、MyBatis-Plus 与 JWT 构建业务服务、权限认证及数据持久化模块；Web 管理端基于 Vue 3、Element Plus 与 ECharts 实现库区管理、设备管理、实时监测、阈值规则、告警中心、工单闭环和数据分析等功能；微信小程序端支持现场员工进行告警查看与工单处理。系统通过 WebSocket 实现设备状态与告警信息的实时推送，并结合阈值规则与告警分级机制构建可追溯的异常处置闭环。同时，引入兼容式大语言模型接口，结合数据库上下文检索机制，将库区、设备、告警及工单等业务数据组织为结构化上下文，通过 deepseek-v4-pro 模型生成风险分析、告警总结及处置建议，从而增强系统的智能分析能力与辅助决策能力。\n\n"
        "系统测试结果表明，本文设计的冷链仓储安全管理系统能够稳定完成实时监测、异常告警、工单流转、数据分析及 AI 智能问答等核心业务流程，具有较好的稳定性、实时性与可扩展性。系统实现了冷链仓储安全管理的数据化、可视化与智能化，验证了大语言模型在冷链仓储安全管理场景中的应用可行性，对智慧物流与智能仓储领域的信息化建设具有一定参考价值。"
    )
    set_text(doc.paragraphs[23], abstract)
    set_text(doc.paragraphs[24], "关键词：冷链仓储；物联网；大语言模型；Spring Boot；告警工单闭环")

    english_abs = (
        "Cold chain warehousing is an important part of ensuring the quality and safety of fresh food, medical vaccines and other temperature-sensitive materials. With the expansion of cold chain logistics, traditional warehousing management has gradually exposed limitations in real-time monitoring, abnormal response, work order closed-loop handling and intelligent data analysis. To address these problems, this thesis designs and implements a cold chain warehousing safety management system based on a large language model. The system aims to improve abnormality discovery efficiency, business collaboration and intelligent analysis capability in cold chain warehousing scenarios.\n\n"
        "The system adopts a front-end and back-end separated architecture. The back-end uses Spring Boot, MySQL, Redis, MyBatis-Plus and JWT to build business services, authentication and data persistence modules. The Web management client uses Vue 3, Element Plus and ECharts to implement warehouse area management, device management, real-time monitoring, threshold rules, alert center, work order closed loop and data analysis. The WeChat mini program supports field staff in viewing alerts and processing work orders. WebSocket is used to push device status and alert information in real time. A traceable abnormal handling loop is constructed through threshold rules, alert grading and work order status flow. In addition, an OpenAI-compatible large model interface is introduced. Business data such as warehouse areas, devices, alerts and work orders are organized as structured context, and the deepseek-v4-pro model is used to generate risk analysis, alert summaries and handling suggestions.\n\n"
        "The test results show that the system can stably complete core business processes such as real-time monitoring, abnormal alerts, work order circulation, data analysis and AI question answering. The system has good stability, real-time performance and extensibility. It realizes data-driven, visual and intelligent cold chain warehousing safety management, and verifies the feasibility of applying large language models to this scenario."
    )
    set_text(doc.paragraphs[29], english_abs)
    set_text(
        doc.paragraphs[30],
        "KEY WORDS: Cold Chain Warehousing; Internet of Things; Large Language Model; Spring Boot; Alert and Work Order Closed Loop",
    )

    set_text_with_cites(
        doc.paragraphs[221],
        [
            "随着物联网、云计算与人工智能技术的快速发展，冷链仓储逐渐向数字化、智能化方向演进。冷链仓储作为保障生鲜食品、医药疫苗及其他温控敏感物资质量安全的重要环节，对温度、湿度、设备运行状态以及异常响应效率具有较高要求。已有研究表明，物联网感知、在线监控和全链路追踪能够有效提升冷链物流过程的可观测性",
            ("[1-3]", "sup"),
            "。在智慧物流与现代供应链体系建设背景下，如何通过信息化手段实现冷链仓储环境的实时监测、风险预警与智能管理，已成为智能仓储领域的重要研究方向。",
        ],
    )
    set_text_with_cites(
        doc.paragraphs[222],
        [
            "传统冷链仓储管理系统虽然能够实现基础数据采集与设备监测，但在异常响应、业务协同及智能分析等方面仍存在不足。一方面，部分系统停留在温湿度展示与简单阈值报警阶段，缺少告警研判、工单流转与结果验收等闭环机制；另一方面，固定图表和静态报表难以对复杂业务数据进行综合分析，难以满足管理人员对风险研判与智能决策的需求",
            ("[4-6]", "sup"),
            "。",
        ],
    )
    set_text(doc.paragraphs[223], "")
    set_text(doc.paragraphs[224], "")
    set_text_with_cites(
        doc.paragraphs[226],
        [
            "大语言模型在自然语言理解、知识归纳与智能分析方面展现出较强能力，为管理系统的数据分析和辅助决策提供了新的技术思路",
            ("[7-9]", "sup"),
            "。将模型能力引入冷链仓储安全管理，不是替代确定性的阈值规则和业务流程，而是在规则判断与人工研判之间提供自然语言解释、风险总结和处置建议，使管理人员能够以更低成本理解多源数据之间的关联。",
        ],
    )
    set_text(
        doc.paragraphs[227],
        "针对上述问题，本文以冷链仓储安全管理为研究对象，设计并实现一套基于大语言模型的冷链仓储安全管理系统。系统以前后端分离架构为基础，通过实时监测、阈值规则、告警分级与工单闭环机制，实现冷链仓储异常的动态感知与全过程管理；同时，引入兼容式大语言模型接口，结合数据库上下文检索机制，对库区、设备、告警与工单等业务数据进行组织与分析，为管理人员提供风险分析、告警总结及处置建议。",
    )
    set_text(
        doc.paragraphs[229],
        "本课题具有一定的理论意义与应用价值。从理论层面看，本文探索了大语言模型与冷链仓储安全管理系统的融合方式，为大语言模型在物联网管理与智慧物流领域中的应用提供参考；同时，通过将实时监测、告警闭环与智能分析相结合，有助于完善智能仓储系统的整体设计思路。",
    )
    set_text(
        doc.paragraphs[230],
        "从应用层面看，本文设计的系统能够提高冷链仓储异常响应效率，优化管理人员与现场员工之间的业务协同流程，并增强数据分析与风险预警能力，对智慧仓储、冷链物流及企业数字化管理具有一定推广价值。该课题涵盖前后端开发、数据库设计、移动端协同以及 AI 智能分析等内容，能够体现软件工程专业毕业设计对综合工程实践能力的要求。",
    )
    set_text_with_cites(
        doc.paragraphs[233],
        [
            "国外冷链物流研究起步较早，研究重点从早期温控设备管理逐步扩展到物联网感知、全链路追踪、食品安全监管和智能决策。相关研究通过 RFID、GPS、无线传感网络和云平台构建货物追踪系统，使运输和仓储过程中的关键环境参数能够被持续记录",
            ("[1-3]", "sup"),
            "。近年来，边缘计算和云原生平台也被用于提高冷链监测系统的实时性和可扩展性。",
        ],
    )
    set_text_with_cites(
        doc.paragraphs[234],
        [
            "在智能分析方面，国外研究更加关注多源数据融合和异常预测。通过对温度曲线、设备负载和历史维护记录进行建模，可以识别冷库设备潜在故障和温控失效趋势。大语言模型出现后，自然语言问答、知识库检索和业务系统数据结合成为新的应用方向",
            ("[7-9]", "sup"),
            "。",
        ],
    )
    set_text_with_cites(
        doc.paragraphs[236],
        [
            "国内冷链物流规模持续扩大，食品、医药和社区零售场景对冷链仓储管理提出了更高要求。相关系统通常包括温湿度采集、设备管理、报警通知、统计报表和权限管理等功能，常结合 Web 管理平台、数据库和移动端入口实现业务协同",
            ("[4-6]", "sup"),
            "。",
        ],
    )
    set_text_with_cites(
        doc.paragraphs[237],
        [
            "在人工智能应用方面，国内系统逐步从固定规则预警向智能分析扩展。部分系统引入机器学习进行异常检测或趋势预测，也有系统利用大语言模型实现智能客服和知识问答。不过，现有实践中仍存在模型与业务数据结合不够紧密、输出缺乏结构化约束、异常处理流程与模型建议脱节等问题",
            ("[7,10]", "sup"),
            "。",
        ],
    )
    set_text(
        doc.paragraphs[239],
        "综合国内外研究与常见工程实践可以发现，传统冷链仓储系统主要存在四类问题。第一，数据展示与业务处置脱节，页面能够看到温湿度数据，但异常研判、派工和验收缺少闭环。第二，系统交互以管理端为主，现场员工移动端功能不足，处理反馈仍依赖线下沟通。第三，数据分析多为固定图表，无法灵活回答管理人员提出的综合问题。第四，系统测试数据不足时，页面、接口和告警流程难以完整验证。",
    )
    set_text(
        doc.paragraphs[240],
        "此外，AI 功能在管理系统中的落地不应停留在简单聊天窗口。若模型无法读取业务数据，回复内容容易变成通用建议；若直接把大量原始数据交给模型，又可能造成上下文冗余和安全风险。因此，系统需要在后端建立受控的数据检索和上下文组织机制，使智能助手能够在可控范围内结合数据库信息完成分析。",
    )
    set_text(doc.paragraphs[241], "")
    set_text(doc.paragraphs[243], "1.3.1 总体目标")
    set_text(
        doc.paragraphs[244],
        "本文的总体目标是形成一个面向冷链仓储安全管理的综合软件平台。系统应支持管理员和仓储管理人员通过 Web 端完成库区、设备、阈值、告警、工单、数据分析和系统管理操作，支持现场员工通过小程序查看告警和处理工单，支持大模型智能助手从数据库上下文中提取关键数据并生成分析建议，从而实现冷链仓储异常发现、处置、验收和复盘的闭环管理。",
    )

    set_text_with_cites(
        doc.paragraphs[258],
        [
            "Spring Boot 是 Java Web 开发中常用的快速开发框架，具有自动配置、起步依赖、内嵌容器和约定优于配置等特点，能够降低应用初始化和依赖整合成本",
            ("[11]", "sup"),
            "。该框架通常与 Spring MVC、Spring Security、数据访问框架和配置管理机制结合使用，适合构建 RESTful 风格的业务服务。",
        ],
    )
    set_text_with_cites(
        doc.paragraphs[260],
        [
            "后端分层设计通常将 Web 控制、业务处理、数据访问和数据模型进行职责拆分。Controller 层负责请求接收与响应封装，Service 层负责业务规则和事务边界，Mapper 或 Repository 层负责数据库访问，Entity、DTO 和 VO 等对象用于隔离持久化结构与接口数据结构",
            ("[11]", "sup"),
            "。",
        ],
    )
    set_text(doc.paragraphs[261], "分层结构有助于降低模块耦合度，提高代码可维护性和可测试性。当业务规则、接口协议或数据库结构发生变化时，合理的分层边界能够减少修改扩散，使系统在后续功能扩展中保持较清晰的组织结构。")
    set_text(doc.paragraphs[262], "")
    set_text(doc.paragraphs[263], "2.1.3 RESTful API 接口风格")
    set_text(doc.paragraphs[264], "RESTful API 以资源为核心组织接口，通过 HTTP 方法表达查询、新增、修改和删除等操作，并使用统一响应结构传递状态码、提示信息和业务数据。该风格具有接口语义清晰、前后端协作方便、便于调试和扩展等特点，是管理系统中常见的接口设计方式。")
    set_text_with_cites(doc.paragraphs[267], ["MySQL 是常见的关系型数据库管理系统，适合保存结构化业务数据和具有明确关联关系的数据对象。关系型数据库通过表、主键、外键、索引和事务机制保证数据组织的规范性与一致性，并支持条件查询、聚合统计和分页检索等典型业务操作", ("[17]", "sup"), "。"])
    set_text(doc.paragraphs[269], "Redis 是基于内存的数据存储系统，支持字符串、哈希、列表、集合和有序集合等多种数据结构，具有读写速度快和过期策略灵活等特点。它常用于缓存热点数据、保存临时状态、实现分布式锁或支撑高频访问场景。")
    set_text_with_cites(doc.paragraphs[271], ["MyBatis-Plus 在 MyBatis 基础上提供通用 CRUD、条件构造器、分页插件和代码生成等能力，能够减少重复 SQL 编写，提高常见数据访问场景的开发效率", ("[16]", "sup"), "。在复杂查询场景中，仍可结合自定义 SQL 保持查询灵活性。"])
    set_text(doc.paragraphs[274], "JWT 是一种常见的无状态认证机制。用户登录成功后，服务端生成包含用户标识、角色信息和过期时间的 Token，客户端在后续请求中通过请求头携带该 Token。服务端解析 Token 并验证签名后，即可判断请求身份。")
    set_text(doc.paragraphs[276], "角色权限模型用于描述不同用户身份对系统资源的访问范围。常见的 RBAC 模型将用户、角色和权限进行关联，通过角色降低直接分配权限带来的管理复杂度。权限控制既可以体现在接口访问层，也可以体现在菜单、按钮和页面路由等前端交互层。")
    set_text(doc.paragraphs[278], "认证与权限控制需要关注密码哈希、Token 过期、非法 Token 拦截、越权访问处理和异常响应一致性。对于前后端分离应用，还需要处理跨域请求、接口白名单和客户端令牌存储等问题。")
    set_text_with_cites(doc.paragraphs[281], ["Vue 3 是用于构建单页应用的前端框架，提供响应式数据绑定、组件化开发、组合式 API 和路由生态等能力", ("[13]", "sup"), "。组件化开发可以将页面拆分为可复用的布局、表单、表格、图表和弹窗模块，提高前端代码组织效率。"])
    set_text_with_cites(doc.paragraphs[283], ["Element Plus 是面向 Vue 3 的桌面端组件库，提供表格、表单、弹窗、抽屉、分页、标签、消息提示等后台管理系统常用组件", ("[14]", "sup"), "。组件库能够统一页面交互风格，减少基础控件重复开发。"])
    set_text_with_cites(doc.paragraphs[285], ["ECharts 是常用的数据可视化库，支持折线图、柱状图、饼图、散点图和仪表盘等多种图表类型，适合展示趋势、分布、占比和指标变化", ("[15]", "sup"), "。合理的数据可视化能够帮助用户更快发现异常模式和业务变化。"])
    set_text(doc.paragraphs[288], "WebSocket 是一种基于 TCP 的全双工通信协议，能够在客户端与服务端之间建立持久连接。与短轮询相比，WebSocket 可以降低频繁请求带来的开销，并在状态变化时由服务端主动向客户端推送消息。")
    set_text(doc.paragraphs[290], "实时推送数据通常包括状态变化、告警事件、任务流转和统计指标刷新等类型。推送消息应保持结构清晰，包含消息类型、业务标识、时间戳和必要的数据字段，以便客户端根据消息类型进行局部更新。")
    set_text(doc.paragraphs[292], "实时连接可能受到网络波动、服务重启和客户端休眠等因素影响。因此，实际应用中通常需要结合心跳检测、断线重连、消息确认和接口补偿查询等机制，保证实时体验和数据一致性。")
    set_text(doc.paragraphs[295], "微信小程序由全局配置、页面 WXML、WXSS、JavaScript 和 JSON 配置文件组成，适合构建轻量级移动端业务应用。小程序通过页面栈、组件、生命周期函数和本地缓存机制组织交互逻辑，并可调用平台提供的网络请求、扫码、消息提示等能力。")
    set_text(doc.paragraphs[297], "移动端现场处置交互强调任务入口清晰、状态提示明确和操作步骤简短。由于手机屏幕空间有限，页面通常需要减少复杂筛选项，突出待处理任务、异常级别、处理要求和提交反馈等高频信息。")
    set_text(doc.paragraphs[298], "2.6.3 小程序网络请求机制")
    set_text(doc.paragraphs[299], "小程序通过网络请求接口与后端服务进行数据交互。为了提高可维护性，通常会对请求地址、请求头、身份令牌、错误提示和超时处理进行统一封装，使不同页面能够以一致方式调用后端接口。")
    set_text_with_cites(doc.paragraphs[302], ["大语言模型基于 Transformer 等深度学习结构，在大规模语料上进行预训练，并通过指令微调等方式增强自然语言理解和生成能力", ("[8-10]", "sup"), "。兼容式模型接口通常以统一的 Chat Completions 协议组织输入消息和输出结果，便于不同模型服务之间进行适配。"])
    set_text(doc.paragraphs[303], "2.7.2 检索增强生成原理")
    set_text_with_cites(doc.paragraphs[304], ["检索增强生成是将外部知识检索与语言模型生成相结合的方法。系统先根据用户问题检索相关文档或结构化数据，再将检索结果作为上下文提供给模型，使模型输出能够结合特定业务数据，而不是仅依赖预训练知识", ("[9]", "sup"), "。"])
    set_text(doc.paragraphs[305], "在检索增强生成过程中，需要控制上下文粒度和数据边界。上下文过少可能导致回答缺乏依据，上下文过多则会增加模型输入长度并引入噪声。因此，常见做法是对检索结果进行筛选、摘要和结构化组织，再交由模型完成解释和总结。")
    set_text(doc.paragraphs[306], "")
    set_text(doc.paragraphs[307], "2.7.3 模型输出与异常处理机制")
    set_text(doc.paragraphs[308], "模型输出需要兼顾可读性、结构化程度和业务可解释性。应用层通常会通过提示词约束输出格式，并对模型服务不可用、密钥配置错误、模型名称不可用、响应超时等异常情况进行兜底处理，避免前端页面长时间等待或显示难以理解的错误信息。")

    for idx, text in {
        487: "5.3 用户认证与权限详细设计与实现",
        498: "5.4 实时监测模块详细设计与实现",
        507: "5.5 告警中心与工单闭环详细设计与实现",
        518: "5.6 AI 智能助手详细设计与实现",
        531: "5.7 微信小程序端详细设计与实现",
        542: "5.8 Web 管理端页面详细设计与实现效果",
        600: "5.9 小程序端详细设计与运行效果",
    }.items():
        set_text(doc.paragraphs[idx], text)
    set_text(doc.paragraphs[489], "登录认证模块的详细设计以“身份校验、令牌签发、路由放行、权限控制”为主线。用户在登录页面输入用户名和密码后，前端调用登录接口；后端查询用户信息并验证密码哈希，验证通过后生成 JWT 返回前端；前端保存 Token 并跳转到 Dashboard，后续请求统一携带 Token。该流程的实现效果如图 5-1 所示。")
    set_text(doc.paragraphs[500], "实时监测模块的详细设计围绕“设备状态采集、阈值判断、状态展示、实时推送”展开。后端接收设备最新数据后写入监测数据表，并根据设备状态和阈值规则更新告警状态；前端通过接口查询和实时推送获取最新结果，在页面中展示温湿度、在线状态、告警状态和所属库区。该模块数据流如图 5-2 所示。")
    set_text(doc.paragraphs[509], "告警中心与工单闭环模块的详细设计以异常处置流程为核心。告警由规则判断触发，当监测数据超过阈值或设备离线时，后端创建告警记录并设置级别；管理人员在告警中心完成确认、忽略、解决或转工单操作；现场员工通过小程序处理工单，最终由管理端验收关闭。告警生成与通知流程如图 5-3 所示。")
    set_text(doc.paragraphs[520], "AI 智能助手模块的详细设计包括问题接收、上下文检索、提示词组织、模型调用和结果渲染五个步骤。后端根据用户问题检索数据库中的库区、设备、告警、工单和统计数据，并将结果整理为模型上下文；模型调用通过兼容式 Chat Completions 协议完成，测试模型为 deepseek-v4-pro。调用流程如图 5-4 所示。")
    set_text(doc.paragraphs[533], "小程序端的详细设计面向现场员工的移动处置场景，页面包括登录、工作台、告警、工单、工单创建、工单详情和个人中心。底部标签栏包含工作台、告警、工单和个人中心四个入口，使员工能够快速进入接警、接单、处理和反馈流程。")
    set_text(doc.paragraphs[544], "Dashboard 页面既是 Web 管理端首页，也是系统运行状态的综合界面。页面详细设计采用顶部概览指标、趋势图、区域状态和异常提醒组合方式，展示设备总数、在线数量、告警数量、待处理工单、区域状态和告警趋势。全局搜索和通知按钮位于顶部布局中，便于管理人员快速定位页面和查看异常提醒，页面实现效果如图 5-6 所示。")
    set_text(doc.paragraphs[550], "库区管理页面采用左侧库区树和右侧详情区域布局。左侧树结构用于表达冷链仓储空间层级，设计时保留完整名称显示，避免节点全部省略影响识别；右侧展示库区基础信息、阈值继承、设备统计和操作按钮，使库区维护与状态查看集中在同一页面，页面实现效果如图 5-7 所示。")

    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)

    doc.save(str(dst))
    print(dst)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
